import pandas as pd

import matplotlib.pyplot as plt

import numpy as np

from matplotlib.ticker import MultipleLocator, FormatStrFormatter, FixedLocator
from datetime import datetime
from windrose import WindroseAxes

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

'''
Исходные данные - файлы arhiv_mcensk_05_21.csv, arhiv_orel_05_21.csv, arhiv_verh_05_21.csv
Результат - в зависимости от выбора файла исходных данных, формируются два результирующих файла
wrose _____.jpg
Роза ветров в ________. График и таблица.docx




Программа предназначена для работы с архивным файлом погоды, скачанным в формате EXCEL 
с сайта rp5.ru 
Предыдущей программой (__________) этот файл преобразован в файл csv, убраны 8 строк текстового заголовка,
переименованы столбцы, сформирован новый столбец precipitation, значения в котором равны 1, если
идет дождь или ливень на дату/время проведения измерения, 2, если идет снег и 0, если осадков не наблюдается.
Столбцы имеют следующие наименования - 
dt_time - дата и время измерения
t_tek - текущая температура
-
-
-
Wind_dir- направление ветра, преобразованное в строковый формат по типу 
    S -южный ветер
    SSW - юго-юго-западный ветер
     ...
    CALM - штиль
wind_speed - скорость ветра в метрах в секунду
-
-
-
...
snow, cm - толщина снежного покрова
precipitetion - наличие дождя (1), снега (2) или отсутствие осадков (0)

При работе программы требуется ввод следующих значений:
наименования файла
начальной даты/времени в формате ДД.ММ.ГГГГ ЧЧ.ММ
после ввода исходных данных формируется датафрейм-срез данных от начальной даты до конечной

Результатом выполнения программы будут по запросу
график температуры
толщина снежного покрова
роза ветров, причем формат розы ветров согласовывается с пользователем.
Например, можно построить розу ветров за ВЕСЬ период, за период наличия осадков в виде снега
при скорости ветра 3 и более метров в секунду (здесь учитывается значение _________, равное 2)
Это так называемая метельная роза ветров
Можно также построить в полярных координатах график преимущественного направления ветра
 (роза ветров без учета скорости ветра, с учетом направления и частоты ветра по этому направлению)  
'''

def pivot_table_to_word(pivot, l_list, region):

    headers = ['С', 'ССВ', 'СВ', 'СВС', 'В', 'ВЮВ', 'ЮВ', 'ЮЮВ', 'Ю', 
               'ЮЮЗ', 'ЮЗ', 'ЗЮЗ', 'З', 'ЗСЗ', 'СЗ', 'ССЗ', 'Всего,%']
    deg = [0.0, 22.5, 45.0, 67.5, 90.0, 112.5, 135.0, 157.5, 180.0, 
           202.5, 225.0, 247.5, 270.0, 292.5, 315.0, 337.5, 'All']
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    current_section = document.sections[-1]
    p = document.add_paragraph('Роза ветров в ')
    p.alignment = 1
    p.add_run(region)

    document.add_picture(l_list[3] + '.jpg', width=Inches(6))
    new_width, new_height = current_section.page_height, \
                            current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    table = document.add_table(rows=1, cols=18)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Напр. ветра/скорость ветра, м/с'

    li_0 = pivot.columns.tolist()

    for j in range(1, 18):
        hdr_cells[j].text = headers[j - 1]
        hdr_cells[j].paragraphs[0].paragraph_format.alignment = \
                    WD_TABLE_ALIGNMENT.CENTER
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(0, pivot.shape[0]):

        row_cells = table.add_row().cells

        l = list(pivot.iloc[i])
        li = list(pivot.index)

        li[len(li) - 1] = 'Всего,%'
        row_cells[0].text = str(li[i])
        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Ввиду того, что в сводной таблице pivot
        # могут отсутствовать некоторые направления ветра, вводится переменная j_sdvig

        j_sdvig = 0
        for j in range(0, 17):

            if li_0[j - j_sdvig] == deg[j]:

                row_cells[j + 1].text = "{:.2f}".format(l[j - j_sdvig])
                row_cells[j + 1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT

            else:

                j_sdvig += 1
                row_cells[j + 1].text = '0.00'
                row_cells[j + 1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT

    try:
        document.save('Роза ветров в ' + l_list[1] + 
                      '. График и таблица.docx')
        file_name = 'Роза ветров в ' + l_list[1] + \
                    '. График и таблица.docx saved'

    except:
        document.save('Роза ветров в ' + l_list[1] + '. График и таблица1.docx')
        file_name = 'Роза ветров в ' + l_list[1] + '. График и таблица1.docx saved'
    return (file_name)


def df_preparation(b_w, snow=True, metel=True):

    b_w['wd'] = b_w.wd.apply(lambda x: np.nan if x == 'CALM' else x)

    if snow:
        # выбор только тех значений, где sn=2
        b_w['sn'] = b_w.sn.apply(lambda x: np.nan if x != 2 else x)

    #        b_w = b_w.dropna()
    # выбор только тех значений, где скорость ветра 3 или более
    if metel:
        b_w['ws'] = b_w.ws.apply(lambda x: np.nan if x < 3 else x)
    #        b_w = b_w.dropna()
    b_w = b_w.dropna()

    return b_w

def wind_rose(b_w, region, r_n, nw=False):

    # nw - признак того, что роза ветров строится по преимущественному направлению, независимо от скорости
    w_dir = {'N': 0, 'NNE': 22.5, 'NE': 45, 'ENE': 67.5, 'E': 90, 'ESE': 112.5,
             'SE': 135, 'SSE': 157.5, 'S': 180,'SSW': 202.5, 'SW': 225, 
             'WSW': 247.5, 'W': 270, 'WNW': 292.5, 'NW': 315, 'NNW': 337.5}
    b_w['wd'] = b_w.wd.apply(lambda x: w_dir[x])
    ax = WindroseAxes.from_ax()
    if nw:
        b_w.ws = 4
        ax.contour(b_w.wd, b_w.ws, bins=np.arange(0, 16, 1), colors='k')
    else:
        if r_n == 0:
            ax.contourf(b_w.wd, b_w.ws, bins=np.arange(0, 16, 1))
        elif r_n == 1:
            ax.contour(b_w.wd, b_w.ws, bins=np.arange(0, 16, 1), colors='k')
        else:
            b_w.ws = 4
            ax.contour(b_w.wd, b_w.ws, bins=np.arange(0, 16, 1), colors='k')
    ax.set_xticklabels(['С', 'СЗ', 'З', 'ЮЗ', 'Ю', 'ЮВ', 'В', 'СВ'])
    ax.set_theta_zero_location('N')
    #    ax.set_legend()
    plt.savefig(region + '.jpg')
    return


def obr_file(archiv, date_n, date_k, snow, metel, r_n):

    time_1 = [' 00:00', ' 03:00', ' 06:00', ' 09:00', ' 12:00', ' 15:00', ' 18:00', ' 21:00']
    file_csv = pd.read_csv(archiv, encoding='cp1251')

    dict = {
        'archiv_mcensk_05_21.csv': ['wrose_met Мценск за 10 лет.jpg',
                                   'Мценске', 'Мценск за период ', 'wrose Мценск'],
        'archiv_orel_05_21.csv': ['wrose_met Орел за 10 лет.jpg', 'Орле',
                                 'Орел за период ', 'wrose Орел'],
        'archiv_verh_05_21.csv': ['wrose_met Верховье за 10 лет.jpg', 'Верховье',
                                 'Верховье за период ',
                                 'wrose Верховье']}



    for i in range(8):
        date_1 = date_n + time_1[i]
        ind = file_csv.index[file_csv['dt_time'] == date_1].tolist()
        if len(ind)!=0:
            date_n = date_1
            break






    index_n = file_csv.loc[file_csv['dt_time'] == date_n].index
    for i in range(7,0,-1):
        date_1 = date_k + time_1[i]
        ind = file_csv.index[file_csv['dt_time'] == date_1].tolist()
        if len(ind)!=0:
            date_k = date_1
            break
    index_k = file_csv.loc[file_csv['dt_time'] == date_k].index




    '''
если нет записи с временем 0:00, то нужно проверить на 03:00, затем на 06.00 
и так далее аналогично для date_k

    получение среза данных от начальной до конечной даты
    Поскольку файл отсортирован по убыванию даты и времени, ..............
    '''
    file_csv_daten_datek = file_csv.loc[index_k[0] - 1:index_n[0], :]

    b_wind = pd.DataFrame()
    b_wind['wd'] = file_csv_daten_datek.Wind_dir
    b_wind['ws'] = file_csv_daten_datek.wind_speed
    b_wind['sn'] = file_csv_daten_datek.precipitation
    b_wind = df_preparation(b_wind, snow, metel)
    b_w = b_wind
    region = dict[archiv][1] + ' ' + date_n + ' ' + date_k
    wind_rose(b_wind, dict[archiv][3], r_n, False)
    res = b_wind.groupby(['ws', 'wd']).size().reset_index(name = 'count')
    deg = [0, 22.5, 45, 67.5, 90, 112.5, 135, 157.5, 180, 202.5, 225, 247.5, 
           270, 292.5, 315, 337.5]
    func = lambda x: round(100 * x.count() / res.shape[0], 2)
    pivot = pd.pivot_table(res, values='count', index=['ws'], 
                           columns=['wd'], aggfunc=func, margins=True)
    pivot = pivot.fillna(0)
    
    if snow:
        sn = 'Осадки в виде снега, '
    else:
        sn = 'Независимо от осадков,'
    if metel:
        sn = sn + 'ветер 3 и более м/с. '
    else:
        sn = sn + 'независимо от скорости ветра.'
    region = dict[archiv][1] + '. Данные с ' + date_n + ' по ' \
                            + date_k + '\n' + sn
    file_name = pivot_table_to_word(pivot, dict[archiv], region)
    return (file_name)
